
# Importação do botcity modo web
from botcity.web import WebBot, Browser, By
# Import for integration with BotCity Maestro SDK
from botcity.maestro import *
# Disable errors if we are not connected to Maestro
BotMaestroSDK.RAISE_NOT_CONNECTED = False
#Importação das opções do navegador
from botcity.web.browsers.chrome import default_options
#atualizar o chromedriver automaticamente
from webdriver_manager.chrome import ChromeDriverManager

from assets import usuario, senha, mover_arquivos, link
import tkinter as tk
from tkinter import simpledialog
from docx import Document
from docx.shared import Inches


class Bot(WebBot):
    
   
    setores = ['Solicitante', 'Desenvolvimento', 'Teste', 'Homologacao', 'Producao']

    documento =  Document()

    #login
    def login(self, user, password):
        self.find_element('exampleInputEmail', By.ID).send_keys(str(user))
        self.wait(1000)
        self.find_element('exampleInputPassword', By.ID).send_keys(str(password))
        self.wait(1000)
        self.find_element('/html/body/div[2]/div/div/div/div/div/div/form/button', By.XPATH).click()

    #Procurar demanada
    def search_demand(self, demand):
        self.wait(1000)
        self.find_element('//*[@id="example_filter"]/label/input', By.XPATH).send_keys(str(demand))
        self.wait(2000)
        self.find_element('//*[@id="item-id"]/a', By.XPATH).click()

    #Extrair informações da guia solicitante
    def extract_info_solicitante(self, i):
        solicitante =  self.find_element('S01', By.ID).text
        gerencia = self.find_element('S02', By.ID).text
        data_criacao = self.find_element('S03', By.ID).text     
        self.get_screenshot(filepath=f'{self.setores[i]}.png')  
        print(f'Solicitante {solicitante}  - Gerencia {gerencia} - Data {data_criacao}')  
         
        p = self.documento.add_paragraph('Solicitantes')
        p = self.documento.add_paragraph(f'Solicitante: {solicitante}')
        p = self.documento.add_paragraph(f'Gerencia do solicitante: {gerencia}')
        p = self.documento.add_paragraph(f'Data de criação: {data_criacao}')
        self.documento.add_picture('Solicitante.png', width=Inches(6.50))

        
    
    #Extrair informações da guia desenvolvimento
    def extract_info_desenvolvimento(self, i):
        desenvolvedor = self.find_element('D01', By.ID).text         
        gerencia = self.find_element('D02', By.ID).text         
        inicio_desenvolvimento = self.find_element('D03', By.ID).text         
        fim_desenvolvimento = self.find_element('D04', By.ID).text         
        liberacao_teste = self.find_element('D06', By.ID).text 
        self.get_screenshot(filepath=f'{self.setores[i]}.png')  
        # print(f'Desenvolvedor {desenvolvedor}  - Gerencia {gerencia} - Data inicio {inicio_desenvolvimento} - Data fim {fim_desenvolvimento} -Liberação pra teste {liberacao_teste} ')
        
        p = self.documento.add_paragraph('Desenvolvimento')
        p = self.documento.add_paragraph(f'Desenvolvedor: {desenvolvedor}')
        p = self.documento.add_paragraph(f'Gerencia do desenvolvedor: {gerencia}')
        p = self.documento.add_paragraph(f'Inicio do desenvolvimento: {inicio_desenvolvimento}')
        p = self.documento.add_paragraph(f'Fim do desenvolvimento: {fim_desenvolvimento}')
        p = self.documento.add_paragraph(f'Data de liberação para o teste: {liberacao_teste}')
        self.documento.add_picture('Desenvolvimento.png', width=Inches(6.50))
        
        
    
    #Extrair informações da guia testes
    def extract_info_testes(self, i):
        testador = self.find_element('T01', By.ID).text 
        gerencia = self.find_element('T02', By.ID).text 
        inicio_teste = self.find_element('T03', By.ID).text         
        fim_teste = self.find_element('T04', By.ID).text    
        liberacao_homologacao = self.find_element('T06', By.ID).text
        self.get_screenshot(filepath=f'{self.setores[i]}.png')  
        # print(f'Testador {testador}  - Gerencia {gerencia} - Data inicio {inicio_teste} - Data fim {fim_teste} -Liberação pra homologacao {liberacao_homologacao} ')  

            
        p = self.documento.add_paragraph('Teste')
        p = self.documento.add_paragraph(f'Testador: {testador}')
        p = self.documento.add_paragraph(f'Gerencia do testador: {gerencia}')
        p = self.documento.add_paragraph(f'Inicio do teste: {inicio_teste}')
        p = self.documento.add_paragraph(f'Fim do teste: {fim_teste}')
        p = self.documento.add_paragraph(f'Data de liberação para o teste: {liberacao_homologacao}')
        self.documento.add_picture('Teste.png', width=Inches(6.50))

        
    #Extrair informações da guia homologação
    def extract_info_homologacao(self, i):
        homologador = self.find_element('H02', By.ID).text
        gerencia = self.find_element('H03', By.ID).text    
        data_homologacao = self.find_element('H04', By.ID).text
        aceito_por = self.find_element('H05', By.ID).text
        data_aceitacao = self.find_element('H06', By.ID).text
        self.get_screenshot(filepath=f'{self.setores[i]}.png')  
        # print(f'Homologador {homologador}  - Gerencia {gerencia} - Homologado em  {data_homologacao} - aceito por {aceito_por} - Data de aceitação {data_aceitacao} ') 
        
        p = self.documento.add_paragraph('Homologação')
        p = self.documento.add_paragraph(f'Homologador: {homologador}')
        p = self.documento.add_paragraph(f'Gerencia do homologador: {gerencia}')
        p = self.documento.add_paragraph(f'Data de homologação: {data_homologacao}')
        p = self.documento.add_paragraph(f'Aceito em: {data_aceitacao}')
        p = self.documento.add_paragraph(f'Aceito por: {aceito_por}')
        self.documento.add_picture('Homologacao.png', width=Inches(6.50))
        
      

    #Extrair iformações da guia produção
    def extract_info_producao(self, i):
        versao_promovida = self.find_element('P01', By.ID).text
        promovido_por = self.find_element('P02', By.ID).text
        gerencia = self.find_element('P03', By.ID).text
        data_promocao =  self.find_element('P04', By.ID).text
        self.get_screenshot(filepath=f'{self.setores[i]}.png')  
        # print(f'Versão {versao_promovida}   - Promovido por  {promovido_por} - Gerencia {gerencia} - Data da promoção {data_promocao} ')
       
        p = self.documento.add_paragraph('Produção')
        p = self.documento.add_paragraph(f'Versão do software: {versao_promovida}')
        p = self.documento.add_paragraph(f'Promovida por: {promovido_por}')
        p = self.documento.add_paragraph(f'Gerencia: {gerencia}')
        p = self.documento.add_paragraph(f'Data de promoção: {data_promocao}')
        self.documento.add_picture('Producao.png', width=Inches(6.50))
        
        

    #Verificar qual a guia correta para cada função
    def extract_info_demand(self):
        abas = ['solic', 'desenv', 'teste', 'homolog', 'prod']
        
        for i, item in enumerate(abas):
            aba = self.find_element(f'//*[@id="{item}"]/a', By.XPATH)

            if 'Solicitante' in aba.text:
                aba.click()
                self.extract_info_solicitante(i)

            elif 'Desenvolvimento' in aba.text:
                aba.click()
                self.extract_info_desenvolvimento(i)

            elif 'Testes' in aba.text:
                aba.click()
                self.extract_info_testes(i)   

            elif 'Homologação/Aceite' in aba.text:
                aba.click()
                self.extract_info_homologacao(i)
                     
            elif 'Produção' in aba.text:
                aba.click()
                self.extract_info_producao(i)


    def main(self, execution=None):
        # Runner passes the server url, the id of the task being executed,
        # the access token and the parameters that this task receives (when applicable).
        maestro = BotMaestroSDK.from_sys_args()
        ## Fetch the BotExecution with details from the task, including parameters
        execution = maestro.get_execution()

        print(f"Task ID is: {execution.task_id}")
        print(f"Task Parameters are: {execution.parameters}")

        # modo oculto
        self.headless = False

        #perfil
        profile_path = r'C:\Users\suporteproc\AppData\Local\Google\Chrome\User Data\Profile Selenium'
       
        #opções do navegador
        def_options = default_options(
            headless=self.headless,
            user_data_dir= profile_path  # Inform the profile path that wants to start the browser
        ) 
        self.options = def_options

        #navegador onde está sendo executado
        self.browser = Browser.CHROME

        #chrome drive
        self.driver_path = ChromeDriverManager().install()

        #abrir o site 
        self.browse(link)
        #Maximizar a janela
        self.maximize_window()

        # Implement here your logic...
        pasta_origem = r'D:\rpa\jornada rpa'

        self.login(usuario, senha)

        root = tk.Tk()
        root.withdraw()
        demanda = simpledialog.askstring("Input", "Informe o numero da demanda:")



        self.search_demand(demanda)
        self.extract_info_demand()

        self.documento.save(f'{demanda}.docx')

        mover_arquivos(pasta_origem, demanda)

        # Wait 3 seconds before closing
        self.wait(3000)

        # Finish and clean up the Web Browser
        # You MUST invoke the stop_browser to avoid
        # leaving instances of the webdriver open
        self.stop_browser()

        # Uncomment to mark this task as finished on BotMaestro
        # maestro.finish_task(
        #     task_id=execution.task_id,
        #     status=AutomationTaskFinishStatus.SUCCESS,
        #     message="Task Finished OK."
        # )

    def not_found(self, label):
        print(f"Element not found: {label}")


if __name__ == '__main__':
    bot = Bot()
    bot.main()
